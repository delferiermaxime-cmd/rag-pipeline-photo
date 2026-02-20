# -*- coding: utf-8 -*-
import asyncio
import logging
import os
import re
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
logger = logging.getLogger(__name__)

# Semaphore : limite à 1 conversion Docling simultanée
# Docling/PyTorch n'est pas thread-safe — plusieurs conversions parallèles
# causent "Cannot copy out of meta tensor" (conflit de device CPU/meta)
# Semaphore initialisé à la demande dans le bon event loop
_DOCLING_SEMAPHORE = None

def _get_semaphore():
    global _DOCLING_SEMAPHORE
    if _DOCLING_SEMAPHORE is None:
        _DOCLING_SEMAPHORE = asyncio.Semaphore(1)
    return _DOCLING_SEMAPHORE
IMAGES_DIR = "/app/images_storage"
try:
    from docling.document_converter import DocumentConverter, PdfFormatOption
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions, TesseractCliOcrOptions
    from docling.backend.pypdfium2_backend import PyPdfiumDocumentBackend
    _DOCLING_OK = True
    logger.info("Docling chargé avec succès")
except Exception as e:
    _DOCLING_OK = False
    logger.error(f"Docling import échoué: {e}")

EXT_TO_FORMAT: Dict[str, Any] = {}
# Extensions qui doivent être converties en .docx avant passage à Docling
# car Docling valide l'extension du fichier temporaire
DOCX_ALIASES = {".dotx", ".doc", ".odt"}

if _DOCLING_OK:
    EXT_TO_FORMAT = {
        ".pdf":  InputFormat.PDF,
        ".docx": InputFormat.DOCX, ".dotx": InputFormat.DOCX, ".doc": InputFormat.DOCX,
        ".pptx": InputFormat.PPTX, ".ppt":  InputFormat.PPTX,
        ".xlsx": InputFormat.XLSX, ".xls":  InputFormat.XLSX,
        ".odt":  InputFormat.DOCX,
        ".html": InputFormat.HTML, ".htm":  InputFormat.HTML,
        ".epub": InputFormat.HTML,
        ".asciidoc": InputFormat.ASCIIDOC, ".adoc": InputFormat.ASCIIDOC,
    }
else:
    EXT_TO_FORMAT = {
        ".txt": "txt", ".md": "txt", ".csv": "txt",
        ".html": "txt", ".htm": "txt",
        ".pdf": "pdf", ".docx": "docx",
    }

def _ensure_images_dir() -> None:
    os.makedirs(IMAGES_DIR, exist_ok=True)

# Sémaphore : limite à 1 conversion Docling simultanée
# Evite les crashes PyTorch quand plusieurs uploads sont lancés en parallèle
_docling_semaphore = None

def _get_semaphore():
    global _docling_semaphore
    if _docling_semaphore is None:
        import asyncio
        _docling_semaphore = asyncio.Semaphore(1)
    return _docling_semaphore


def _build_converter(ext: str) -> "DocumentConverter":
    opts = PdfPipelineOptions()
    opts.do_ocr = True
    opts.do_table_structure = True
    opts.images_scale = 2.0
    opts.generate_page_images = True
    opts.generate_picture_images = True
    # Force Tesseract — évite EasyOCR qui crash sans GPU dans le container
    try:
        opts.ocr_options = TesseractCliOcrOptions(lang=["fra", "eng"])
    except Exception as e:
        logger.warning(f"TesseractCliOcrOptions non disponible: {e}")
    if ext == ".pdf":
        return DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(
                    pipeline_options=opts,
                    backend=PyPdfiumDocumentBackend,
                )
            }
        )
    return DocumentConverter()

def _chunk_markdown(
    markdown: str,
    filename: str,
    max_chars: int = 3000,
    overlap: int = 450,
) -> List[Dict[str, Any]]:
    heading_re = re.compile(r"^(#{1,6}\s.+)$", re.MULTILINE)
    parts = heading_re.split(markdown)
    sections: List[Dict[str, str]] = []
    current_heading = Path(filename).stem
    buffer = ""
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if heading_re.match(part):
            if buffer.strip():
                sections.append({"heading": current_heading, "content": buffer.strip()})
            current_heading = part.lstrip("#").strip()
            buffer = part + "\n"
        else:
            buffer += part + "\n"
    if buffer.strip():
        sections.append({"heading": current_heading, "content": buffer.strip()})
    if not sections:
        sections = [{"heading": Path(filename).stem, "content": markdown}]
    chunks: List[Dict[str, Any]] = []
    last_chunk_tail = ""
    def _add_chunk(title: str, text: str, page: int) -> None:
        nonlocal last_chunk_tail
        prefixed = (last_chunk_tail + "\n\n" + text).strip() if last_chunk_tail else text.strip()
        real_title = title
        for line in prefixed.split("\n"):
            m = re.match(r"^#{1,6}\s+(.+)", line.strip())
            if m:
                real_title = f"{Path(filename).stem} — {m.group(1).strip()}"
                break
        chunks.append({
            "page": page,
            "title": real_title,
            "content": prefixed,
            "chunk_index": len(chunks),
        })
        last_chunk_tail = text.strip()[-overlap:] if len(text.strip()) > overlap else text.strip()
    for i, section in enumerate(sections):
        text = section["content"]
        title = f"{Path(filename).stem} — {section['heading']}"
        page = i + 1
        if len(text) <= max_chars:
            _add_chunk(title, text, page)
        else:
            paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
            buf = ""
            for para in paragraphs:
                if buf and len(buf) + len(para) + 2 > max_chars:
                    _add_chunk(title, buf.strip(), page)
                    buf = para + "\n\n"
                else:
                    buf += para + "\n\n"
            if buf.strip():
                _add_chunk(title, buf.strip(), page)
    return chunks or [{"page": 1, "title": filename, "content": "(vide)", "chunk_index": 0}]

def _save_images_sync(
    result: Any,
    document_id: str,
) -> List[Dict[str, Any]]:
    _ensure_images_dir()
    saved: List[Dict[str, Any]] = []
    try:
        doc = result.document
        for page_no, page in enumerate(doc.pages, start=1):
            if hasattr(page, 'image') and page.image is not None:
                img = page.image
                if hasattr(img, 'pil_image') and img.pil_image is not None:
                    fname = f"{document_id}_page_{page_no}.png"
                    fpath = os.path.join(IMAGES_DIR, fname)
                    img.pil_image.save(fpath, "PNG")
                    saved.append({"page": page_no, "filename": fname, "type": "page"})
        for elem_idx, element in enumerate(doc.elements or []):
            if hasattr(element, 'image') and element.image is not None:
                img = element.image
                if hasattr(img, 'pil_image') and img.pil_image is not None:
                    page_no = getattr(getattr(element, 'prov', [None])[0], 'page', 1) if getattr(element, 'prov', None) else 1
                    fname = f"{document_id}_img_{elem_idx}_p{page_no}.png"
                    fpath = os.path.join(IMAGES_DIR, fname)
                    img.pil_image.save(fpath, "PNG")
                    saved.append({"page": page_no, "filename": fname, "type": "inline"})
    except Exception as e:
        logger.warning(f"Extraction images partielle : {e}")
    logger.info(f"[Images] {len(saved)} images extraites pour document {document_id}")
    return saved

def _convert_sync(
    tmp_path: str,
    ext: str,
    filename: str,
    document_id: str,
) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    """Retourne (chunks, images)"""
    converter = _build_converter(ext)
    result = converter.convert(tmp_path)
    markdown: str = result.document.export_to_markdown()
    if not markdown or not markdown.strip():
        return [{"page": 1, "title": filename, "content": "(document vide)", "chunk_index": 0}], []
    chunks = _chunk_markdown(markdown, filename)
    images = _save_images_sync(result, document_id)
    return chunks, images

def _fallback_parse(file_bytes: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext in (".txt", ".md", ".csv", ".html", ".htm"):
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                return file_bytes.decode(enc)
            except UnicodeDecodeError:
                continue
        return file_bytes.decode("utf-8", errors="replace")
    if ext == ".pdf":
        try:
            import pypdf, io
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            pages = [p.extract_text() for p in reader.pages if p.extract_text()]
            return "\n\n".join(pages) if pages else "(PDF vide ou non lisible)"
        except Exception:
            try:
                import pypdfium2 as pdfium
                pdf = pdfium.PdfDocument(file_bytes)
                return "\n\n".join(pdf[i].get_textpage().get_text_range() for i in range(len(pdf)))
            except Exception as e2:
                raise RuntimeError(f"Impossible de lire le PDF: {e2}")
    if ext in (".docx", ".dotx", ".doc", ".odt"):
        try:
            import docx, io
            doc = docx.Document(io.BytesIO(file_bytes))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            raise RuntimeError(f"Impossible de lire le DOCX: {e}")
    raise ValueError(f"Format non supporté: {ext}")

async def convert_document(
    file_bytes: bytes,
    filename: str,
    document_id: str = "",
) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    """
    Retourne (chunks, images).
    FIX .dotx/.doc : on renomme le fichier temporaire en .docx avant de le passer à Docling
    car Docling valide l'extension du fichier temporaire.
    """
    ext = Path(filename).suffix.lower()

    # Détermine l'extension réelle à utiliser pour le fichier temporaire
    tmp_ext = ext
    if ext in DOCX_ALIASES:
        tmp_ext = ".docx"
        logger.info(f"[Docling] Extension '{ext}' → renommée en '.docx' pour compatibilité Docling")

    if _DOCLING_OK and ext in EXT_TO_FORMAT:
        logger.info(f"[Docling] Conversion '{filename}' ({len(file_bytes):,} bytes)")
        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix=tmp_ext, delete=False, dir="/tmp") as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            # Sémaphore : 1 seule conversion Docling à la fois (évite crash PyTorch concurrent)
            async with _get_semaphore():
                logger.info(f"[Docling] Semaphore acquis pour '{filename}'")
                chunks, images = await asyncio.to_thread(_convert_sync, tmp_path, tmp_ext, filename, document_id)
            logger.info(f"[Docling] {len(chunks)} chunks, {len(images)} images pour '{filename}'")
            return chunks, images
        except Exception as e:
            logger.error(f"[Docling] Erreur '{filename}': {e}", exc_info=True)
            raise RuntimeError(f"Erreur Docling: {e}")
        finally:
            if tmp_path and os.path.exists(tmp_path):
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass
    else:
        logger.info(f"[Fallback] Conversion '{filename}'")
        text = _fallback_parse(file_bytes, filename)
        chunks = _chunk_markdown(text, filename)
        logger.info(f"[Fallback] {len(chunks)} chunks pour '{filename}'")
        return chunks, []
